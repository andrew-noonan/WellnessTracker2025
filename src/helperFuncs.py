# helperFuncs.py

import openpyxl
from docx import Document
from docxtpl import DocxTemplate
import os
from os.path import basename, join
from helperClasses import Question, CheckboxQuestion, DynamicVar

# HELPER FUNCTIONS FOR READING IN QUESTIONS
def parse_conditions(*conditions):
    conditions_list = []
    for condition in conditions:
        if condition:
            conditions_list.append([c.strip() for c in condition.split(',')])
        else:
            conditions_list.append([])
    return conditions_list

def read_sheet(ws, question_class, has_options=False):
    questions = []
    for row in ws.iter_rows(min_row=2, values_only=True):  # Skipping the header row
        question_id, prompt, cond1, cond2, cond3, *_ = row
        question_type = question_id[:2]
        if has_options:
            options = [option.strip() for option in row[5].split(',')] if row[5] else []
            question = CheckboxQuestion(question_id, question_type, cond1, cond2, cond3, prompt, options)
        else:
            question = Question(question_id, question_type, cond1, cond2, cond3, prompt)
        questions.append(question)
    return questions


def importQuestions(filepath):
    workbook = openpyxl.load_workbook(filepath)
    
    # Create a single list to hold all questions
    all_questions = []

    # Read questions from each sheet and extend the all_questions list
    all_questions.extend(read_sheet(workbook['CheckboxQuestions'], CheckboxQuestion, has_options=True))
    all_questions.extend(read_sheet(workbook['YesNoQuestions'], Question))
    all_questions.extend(read_sheet(workbook['Free Response'], Question))
    
    return all_questions

# Example usage
# all_questions = importQuestions('path_to_your_excel_file.xlsx')

# HELPER FUNCTIONS FOR CONDITION EVALUATION
def evaluate_condition(subcondition, questions_dict):
        """Evaluate a single subcondition against the provided dictionary of questions, 
        where question.values is now a list of strings."""
        identifier, expected = subcondition
        question = questions_dict.get(identifier)
        
        # If the question doesn't exist or its values list is empty, return False
        if not question or not question.values:
            return False
        
        # If expected starts with '!', check for absence
        if expected.startswith('!'):
            # Remove the leading '!' for comparison
            expected_value = expected[1:]
            # Return True if the expected value is NOT in the question's values
            return expected_value not in question.values
        else:
            # Return True if the expected value IS in the question's values
            return expected in question.values

def check_question_conditions(question, questions_dict):
        """Updates the conditionsSatisfied attribute for the provided question more efficiently."""
        # Consolidate conditions into a list for easier iteration
        conditions = [question.conditions1, question.conditions2, question.conditions3]
        
        # Check if all condition containers are empty, implying no conditions and thus always satisfied
        if all(not c for c in conditions):
            question.conditionsSatisfied = True
            return
        
        # Otherwise, evaluate each condition
        question.conditionsSatisfied = any(
            all(evaluate_condition(subcondition, questions_dict) for subcondition in condition)
            for condition in conditions if condition
        )

def check_variable_conditions(variable, questions_dict):
        """Updates the conditionsSatisfied attribute for the provided question more efficiently."""
        # Consolidate conditions into a list for easier iteration
        conditions = [variable.conditions1, variable.conditions2, variable.conditions3, variable.conditions4, variable.conditions5, variable.conditions6]
        
        # Check if all condition containers are empty, implying no conditions and thus always satisfied
        if all(not c for c in conditions):
            variable.value = "Err - No conditions"
            return
        
        for index, condition in enumerate(conditions):
             # Evaluate all subconditions for a single condition
            for subcondition in condition:
                if evaluate_condition(subcondition, questions_dict):
                    # If all subconditions are met, set the variable's value to the corresponding value
                    variable.value = variable.allVals[index]

def addDependents(questions_dict):
    for q in questions_dict.values():
        for cond in (q.conditions1, q.conditions2, q.conditions3):
            for subcond in cond:
                questions_dict[subcond[0]].dependents = questions_dict[subcond[0]].dependents | {q.question_id}
# HELPER FUNCTIONS FOR READING AND WRITING WORD DOC

def import_vars_from_template(filepath):
    doc = Document(filepath)
    variables = {}

    # Function to search for variables in a text and update the variables dictionary
    def search_vars_in_text(text, container_label, position, variables_dict):
        start = text.find('{{')
        end = text.find('}}', start + 2)

        while start != -1 and end != -1 and start < end:
            var_name = text[start + 2:end]
            if var_name not in variables_dict:
                variables_dict[var_name] = []

            variables_dict[var_name].append({
                'text': text[start:end+2], 
                'start': start, 
                'end': end + 2,
                'container': container_label,
                'position': position
            })

            start = text.find('{{', end + 2)
            end = text.find('}}', start + 2)

    # Process the document body
    for i, paragraph in enumerate(doc.paragraphs):
        full_text = "".join(run.text for run in paragraph.runs)
        search_vars_in_text(full_text, 'body', i, variables)

    # Process the header(s)
    for i, section in enumerate(doc.sections):
        for header in section.header.paragraphs:
            full_text = "".join(run.text for run in header.runs)
            search_vars_in_text(full_text, f'header_section_{i}', i, variables)

    return variables

def load_static_variables(excel_path):
    workbook = openpyxl.load_workbook(excel_path)
    sheet = workbook['StaticVariables']
    static_vars = {row[0]: row[1] for row in sheet.iter_rows(min_row=2, max_row=5, values_only=True)}
    return static_vars

def load_dynamic_variables(file_path, questions_dict):
    wb = openpyxl.load_workbook(file_path)
    sheet = wb["Linking"]
    variables_dict = {}

    for row in sheet.iter_rows(min_row=2, values_only=True):  # Assuming the first row is headers
        variable_id, conditions1, value1, conditions2, value2, conditions3, value3, conditions4, value4, conditions5, value5, conditions6, value6 = row
        allVals = [v for v in [value1, value2, value3, value4, value5, value6] if v is not None]

        variable = DynamicVar(
            variable_id=variable_id,
            conditions1=conditions1,
            conditions2=conditions2,
            conditions3=conditions3,
            conditions4=conditions4,
            conditions5=conditions5,
            conditions6=conditions6,
            allVals=allVals,
            value='N/A'
        )
        check_variable_conditions(variable, questions_dict)
        variables_dict[variable_id] = variable

    return variables_dict


def replace_template(docx_template_path, context, output_dir):
    """
    Fills in a Word document template with static variables.
    
    Parameters:
    - docx_template_path: Path to the Word document template.
    - context: Dictionary with static variables to fill in the template.
    - output_dir: Directory where the filled-in document will be saved.
    """
    # Load static variables and fill in the template
    doc = DocxTemplate(docx_template_path)
    doc.render(context)
    
    # Construct the output path using the output directory and the original filename
    original_filename = basename(docx_template_path)
    output_path = join(output_dir, original_filename.replace('.docx', '_filled.docx'))
    
    # Save the filled-in document
    doc.save(output_path)
    print(f"Filled-in document saved to: {output_path}")
    return output_path

def finishing_touches(filepath, questions_dict):
    found_variables = import_vars_from_template(filepath)
    variable_values = {}
    #Assign section and subsection numbers
    sec_counter = 0  # Counter for sections
    ssec_counter = 0  # Counter for subsections within the current section
    curSec = 'ZZ'
    for var in sorted(found_variables.keys()):
        if var.startswith("sec"):
            if var[4] != curSec: # New Section
                sec_counter += 1
                ssec_counter = 0  # Reset subsection counter
                curSec = var[4]
                variable_values[var] = f"{sec_counter}.00"
            else:
                ssec_counter += 1
                variable_values[var] = f"{sec_counter}.{ssec_counter:02}"
        else: # Remaining variable is not sec or ssec
            variable_values[var] = format_list_with_oxford_comma(questions_dict[var].values)
    
    # Fill in the template with the section numbers
    doc = DocxTemplate(filepath)
    doc.render(variable_values)
    # Save the filled-in document
    output_path = filepath.replace('.docx', '_final.docx')
    doc.save(output_path)
    print(f"Final document saved to: {output_path}")
    return output_path

def format_list_with_oxford_comma(items):
    # Check the number of items in the list
    num_items = len(items)
    
    if num_items == 0:
        return ''
    elif num_items == 1:
        # Just return the single item
        return items[0]
    elif num_items == 2:
        # Return the two items separated by "and"
        return f"{items[0]} and {items[1]}"
    else:
        # Use an oxford comma for 3 or more items
        return f"{', '.join(items[:-1])}, and {items[-1]}"

def recursive_ask(question, questions_dict):
    if not question.asked and question.conditionsSatisfied:
        # Simulate asking the question (e.g., print or input)
        print(question.prompt)
        response = input('Answer: ')  # Simulate getting an answer
        question.values.extend([r.strip() for r in response.split(',')])
        question.asked = True
        
        for dependent in question.dependents: 
            check_question_conditions(questions_dict[dependent], questions_dict)

        # Recursively ask dependent questions
        for dependent in question.dependents:
            recursive_ask(questions_dict[dependent], questions_dict)

